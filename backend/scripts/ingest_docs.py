"""Ingest PDF/DOCX from uploads/knowledge into PGVector (idempotent by file hash).

Usage (from backend/ or inside the container):
    python scripts/ingest_docs.py
"""

from __future__ import annotations

import hashlib
import os
import sys
from pathlib import Path

# Default FastEmbed cache as specified; Docker compose may override via env.
os.environ.setdefault("FASTEMBED_CACHE_DIR", "./.cache/fastembed")
Path(os.environ["FASTEMBED_CACHE_DIR"]).mkdir(parents=True, exist_ok=True)

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from docx import Document as DocxDocument  # noqa: E402
from langchain_core.documents import Document  # noqa: E402
from langchain_text_splitters import RecursiveCharacterTextSplitter  # noqa: E402
from sqlalchemy import create_engine, text  # noqa: E402

from app.config import get_settings  # noqa: E402
from app.database import init_db  # noqa: E402
from app.services import rag_common  # noqa: E402
from app.services.rag_common import get_vector_store  # noqa: E402

ALLOWED_EXTENSIONS = {".pdf", ".docx"}
CATEGORY_HINTS = (
    (("устав", "регламент", "политик"), "GENERAL"),
    (("безопас", "охран", "safety"), "SAFETY"),
    (("кредит", "займ", "credit"), "CREDIT"),
    (("приказ", "кадр", "hr", "отпуск"), "HR"),
)


def _guess_category(filename: str) -> str:
    name = filename.lower()
    for needles, category in CATEGORY_HINTS:
        if any(n in name for n in needles):
            return category
    return "GENERAL"


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as fh:
        for chunk in iter(lambda: fh.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _extract_text(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".pdf":
        import fitz  # PyMuPDF

        parts: list[str] = []
        with fitz.open(path) as pdf:
            for page in pdf:
                parts.append(page.get_text())
        return "\n".join(parts).strip()
    if ext == ".docx":
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    raise ValueError(f"Unsupported extension: {ext}")


def _pg_url() -> str:
    settings = get_settings()
    url = settings.database_url
    if url.startswith("postgresql://"):
        return url.replace("postgresql://", "postgresql+psycopg://", 1)
    if url.startswith("postgres://"):
        return url.replace("postgres://", "postgresql+psycopg://", 1)
    return url


def _hash_already_ingested(file_hash: str) -> bool:
    """Skip when any embedding already stores this file_hash in cmetadata."""
    engine = create_engine(_pg_url())
    try:
        with engine.connect() as conn:
            row = conn.execute(
                text(
                    """
                    SELECT 1
                    FROM langchain_pg_embedding
                    WHERE cmetadata->>'file_hash' = :h
                    LIMIT 1
                    """
                ),
                {"h": file_hash},
            ).fetchone()
            return row is not None
    except Exception:
        return False
    finally:
        engine.dispose()


def _list_knowledge_files(root: Path) -> list[Path]:
    if not root.exists():
        return []
    return [
        p
        for p in sorted(root.rglob("*"))
        if p.is_file() and p.suffix.lower() in ALLOWED_EXTENSIONS
    ]


def main() -> int:
    get_settings.cache_clear()
    rag_common.get_embeddings.cache_clear()
    rag_common.get_vector_store.cache_clear()

    settings = get_settings()
    init_db()

    knowledge_dir = Path(settings.rag_documents_dir)
    files = _list_knowledge_files(knowledge_dir)
    total = len(files)
    if total == 0:
        print(f"No .pdf/.docx files found in {knowledge_dir.resolve()}")
        return 0

    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=100)
    store = get_vector_store()

    ingested_files = 0
    skipped_files = 0
    total_chunks = 0

    for index, path in enumerate(files, start=1):
        print(f"Processing file {index} of {total}: {path.name}")
        file_hash = _sha256_file(path)
        if _hash_already_ingested(file_hash):
            print(f"  -> skip (hash {file_hash[:12]}... already in PGVector)")
            skipped_files += 1
            continue

        try:
            text_content = _extract_text(path)
        except Exception as exc:
            print(f"  -> ERROR extracting text: {exc}")
            continue

        if not text_content:
            print("  -> skip (empty document)")
            skipped_files += 1
            continue

        category = _guess_category(path.name)
        raw = Document(
            page_content=text_content,
            metadata={
                "source_file": path.name,
                "category": category,
                "file_hash": file_hash,
                "file_path": str(path),
            },
        )
        chunks = splitter.split_documents([raw])
        for chunk_index, chunk in enumerate(chunks):
            chunk.metadata.update(
                {
                    "source_file": path.name,
                    "category": category,
                    "chunk_index": chunk_index,
                    "file_hash": file_hash,
                }
            )

        store.add_documents(chunks)
        ingested_files += 1
        total_chunks += len(chunks)
        print(f"  -> upserted {len(chunks)} chunks (category={category})")

    print(
        f"Done. files_ingested={ingested_files}, files_skipped={skipped_files}, "
        f"chunks={total_chunks}, dir={knowledge_dir}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
