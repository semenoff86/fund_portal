"""Document ingestion pipeline for RAG (DOCX/PDF → PGVector) with versioning."""

from __future__ import annotations

import logging
from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sqlalchemy import func
from sqlalchemy.orm import Session

from app.config import get_settings
from app.models import DocumentStatus, OrderDocument
from app.services import rag_common
from app.services.rag_common import get_vector_store

logger = logging.getLogger(__name__)
settings = get_settings()

ALLOWED_EXTENSIONS = {".pdf", ".docx"}


def _guess_category(filename: str, parent_dir: str) -> str:
    name_lower = filename.lower()
    parent_lower = Path(parent_dir).name.lower()
    if parent_lower not in ("knowledge", "uploads", "documents", ""):
        return parent_lower.upper()
    if "устав" in name_lower:
        return "GENERAL"
    if "безопас" in name_lower or "охран" in name_lower:
        return "SAFETY"
    if "кредит" in name_lower or "займ" in name_lower:
        return "CREDIT"
    if "приказ" in name_lower or "hr" in name_lower or "кадр" in name_lower:
        return "HR"
    return "GENERAL"


def _load_pdf_text(path: Path) -> str:
    import fitz  # PyMuPDF

    text_parts: list[str] = []
    with fitz.open(path) as pdf:
        for page in pdf:
            text_parts.append(page.get_text())
    return "\n".join(text_parts).strip()


def _load_docx_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()


def extract_text_from_file(file_path: str | Path) -> str:
    """Extract plain text from a PDF or DOCX file."""
    path = Path(file_path)
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _load_pdf_text(path)
    if ext == ".docx":
        return _load_docx_text(path)
    raise ValueError(f"Unsupported file type: {ext}")


def load_file_as_document(path: Path) -> Document | None:
    ext = path.suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        return None
    try:
        text = extract_text_from_file(path)
    except Exception as exc:
        logger.warning("Failed to load %s: %s", path, exc)
        return None
    if not text:
        return None
    return Document(
        page_content=text,
        metadata={
            "source_file": path.name,
            "category": _guess_category(path.name, str(path.parent)),
            "file_path": str(path),
        },
    )


def load_documents_from_directory(directory: Path | None = None) -> list[Document]:
    root = directory or Path(settings.rag_documents_dir)
    if not root.exists():
        logger.warning("RAG documents directory not found: %s", root)
        return []
    docs: list[Document] = []
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in ALLOWED_EXTENSIONS:
            doc = load_file_as_document(path)
            if doc:
                docs.append(doc)
    return docs


def split_documents(documents: list[Document]) -> list[Document]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=settings.rag_chunk_size,
        chunk_overlap=settings.rag_chunk_overlap,
    )
    chunks = splitter.split_documents(documents)
    for idx, chunk in enumerate(chunks):
        source = chunk.metadata.get("source_file", "unknown")
        chunk.metadata["chunk_index"] = idx
        chunk.metadata["source_file"] = source
    return chunks


def handle_document_versioning(
    db: Session,
    *,
    filename: str,
    category: str,
    title: str | None = None,
) -> int:
    """
    Archive previous active versions of the same document and return the next version number.

    Matching strategy (first hit wins):
      1. Active docs with the same title (case-insensitive) and category
      2. Active docs whose stored file_path ends with the original filename
    """
    from app.models import DocumentCategory

    try:
        cat = DocumentCategory(category) if not isinstance(category, DocumentCategory) else category
    except ValueError:
        cat = None

    match_title = (title or Path(filename).stem).strip()
    q = db.query(OrderDocument).filter(OrderDocument.is_active.is_(True))
    if cat is not None:
        q = q.filter(OrderDocument.category == cat)

    candidates = q.filter(
        func.lower(OrderDocument.title) == match_title.lower()
    ).all()

    if not candidates and filename:
        # Fallback: match by original filename substring in file_path
        candidates = (
            q.filter(OrderDocument.file_path.ilike(f"%{filename}"))
            .all()
        )

    if not candidates:
        return 1

    max_version = max(d.version for d in candidates)
    for doc in candidates:
        doc.is_active = False
        doc.status = DocumentStatus.ARCHIVED

    db.flush()
    return max_version + 1


def ingest_document(
    file_path: str,
    category: str,
    uploaded_by: int,
    *,
    version: int = 1,
    source_file: str | None = None,
    document_id: int | None = None,
) -> dict:
    """
    Ingest a single PDF/DOCX into the LangChain PGVector collection.

    Steps:
      1. Extract text (PyMuPDF / python-docx)
      2. Split with RecursiveCharacterTextSplitter (chunk_size=800, overlap=100)
      3. Embed via FastEmbed (BAAI/bge-m3)
      4. Upsert chunks with metadata: source_file, category, version, chunk_index, uploaded_by
    """
    path = Path(file_path)
    if not path.exists():
        # Support portal paths like /uploads/knowledge/xxx.pdf
        alt = Path(file_path.lstrip("/"))
        if not alt.exists() and str(file_path).startswith("/uploads/"):
            alt = Path("uploads") / file_path.removeprefix("/uploads/")
        path = alt if alt.exists() else path

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    text = extract_text_from_file(path)
    if not text:
        return {"chunks": 0, "content_text": "", "message": "Empty document"}

    display_name = source_file or path.name
    raw = Document(
        page_content=text,
        metadata={
            "source_file": display_name,
            "category": category,
            "version": version,
            "uploaded_by": uploaded_by,
            "file_path": str(path),
            "document_id": document_id,
        },
    )
    chunks = split_documents([raw])
    for idx, chunk in enumerate(chunks):
        chunk.metadata.update(
            {
                "source_file": display_name,
                "category": category,
                "version": version,
                "chunk_index": idx,
                "uploaded_by": uploaded_by,
                "document_id": document_id,
            }
        )

    store = get_vector_store()
    store.add_documents(chunks)
    return {
        "chunks": len(chunks),
        "content_text": text,
        "message": f"Ingested {len(chunks)} chunks from {display_name}",
    }


def ingest_directory(directory: Path | None = None, *, reset: bool = True) -> dict:
    """
    Load, chunk, embed and upsert documents into PGVector.
    If reset=True, clears the collection before re-ingestion.
    """
    raw_docs = load_documents_from_directory(directory)
    if not raw_docs:
        return {"files": 0, "chunks": 0, "message": "No documents found"}

    chunks = split_documents(raw_docs)
    store = get_vector_store()

    if reset:
        try:
            store.delete_collection()
        except Exception as exc:
            logger.warning("Could not reset collection: %s", exc)
        rag_common.get_vector_store.cache_clear()
        store = get_vector_store()

    store.add_documents(chunks)
    return {
        "files": len(raw_docs),
        "chunks": len(chunks),
        "message": f"Ingested {len(chunks)} chunks from {len(raw_docs)} files",
    }
